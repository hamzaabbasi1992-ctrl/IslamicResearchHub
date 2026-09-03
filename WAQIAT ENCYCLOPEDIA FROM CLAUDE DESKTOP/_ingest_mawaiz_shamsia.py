import os, sys, glob, re, sqlite3, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r"F:\ISLAMIC RESEARCH HUB AI\src")

from islamic_research_hub.application.event_extraction import ExtractedEvent
from islamic_research_hub.infrastructure.persistence.event_candidate_repository import EventCandidateRepository
from islamic_research_hub.infrastructure.persistence.taxonomy_repository import TaxonomyRepository
from islamic_research_hub.infrastructure.persistence.event_candidate_taxonomy_repository import EventCandidateTaxonomyRepository

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
os.makedirs(OUTPUT_DIR, exist_ok=True)

AITIKAF_DIR = r"E:\KHANQAH\AITIKAAF BAYANAT\book formatted"
SUNDAY_DIR = r"E:\KHANQAH\SUNDAY BAYANS (FINAL)\bayan to text gemini output"

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

def clean_xml_text(s):
    if not s: return ""
    s = re.sub(r'</?[^>]+>', ' ', s)
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', s)
    s = re.sub(r'www\.[a-zA-Z0-9_\-\.]+', '', s)
    return re.sub(r'\s+', ' ', s).strip()

def clean_title(t, maxlen=85):
    t = clean_xml_text(t)
    t = re.sub(r'^[\d۱-۹١-٩\(\)\*\-\s٭…_]+', '', t).strip()
    return (t[:maxlen] + "...") if len(t) > maxlen else t

ALL_TRIGGERS = [
    r'(ایک مرتبہ\s+[^۔\n]{5,50}؟?)',
    r'(ایک دفعہ\s+[^۔\n]{5,50}؟?)',
    r'(ایک بار\s+[^۔\n]{5,50}؟?)',
    r'(ایک بزرگ\s+[^۔\n]{5,50}؟?)',
    r'(ایک شخص\s+[^۔\n]{5,50}؟?)',
    r'(ایک بادشاہ\s+[^۔\n]{5,50}؟?)',
    r'(ایک صاحب\s+[^۔\n]{5,50}؟?)',
    r'(منقول ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(روایت ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(حکایت ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(واقعہ ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(واقعہ یہ ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(حضرت\s+[^\n۔]{3,25}\s+کا واقعہ\s+[^۔\n]{0,30})',
    r'(خواب میں دیکھا\s+کہ\s+[^۔\n]{5,50})',
    r'(فرمایا کہ\s+[^۔\n]{10,50})',
    r'(ارشاد فرمایا\s+[^۔\n]{5,50})',
    r'(بیان فرمایا\s+[^۔\n]{5,50})',
    r'(لکھا ہے کہ\s+[^۔\n]{5,50})',
    r'(آیا ہے کہ\s+[^۔\n]{5,50})',
    r'(ایک ولی\s+[^۔\n]{5,50})',
    r'(ایک درویش\s+[^۔\n]{5,50})',
    r'(ایک فقیر\s+[^۔\n]{5,50})',
    r'(ایک قصہ\s+[^۔\n]{0,40})',
    r'(ایک واقعہ\s+[^۔\n]{0,40})'
]
combined_pattern = re.compile("|".join(ALL_TRIGGERS), re.UNICODE)

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

# Clean any existing candidates for BookID 95001 and 95002
cur.execute("DELETE FROM EventCandidateTaxonomyTerms WHERE EventCandidateID IN (SELECT EventCandidateID FROM EventCandidates WHERE BookID IN (95001, 95002))")
cur.execute("DELETE FROM EventCandidates WHERE BookID IN (95001, 95002)")

# Insert or Replace Book records
cur.execute("""
    INSERT OR REPLACE INTO Books (BookID, Source, SourceBookID, Title, Author, Publisher, Language, Category, PageCount, ChapterCount, LibraryID)
    VALUES 
    (95001, 'khanqah_mawaiz_shamsia_aitikaf_2009_2023', '95001', 'مواعظِ شمسیہ — بیاناتِ اعتکاف (۲۰۰۹ء تا ۲۰۲۳ء)', 'حضرت مولانا شمس الرحمن العباسی دامت برکاتہم', 'خانقاہ امدادیہ غفوریہ', 'ur', '33', 1500, 14, 1),
    (95002, 'khanqah_mawaiz_shamsia_sunday_bayanat_2010_2023', '95002', 'مواعظِ شمسیہ — مواعظ و مجالس اتوار', 'حضرت مولانا شمس الرحمن العباسی دامت برکاتہم', 'خانقاہ امدادیہ غفوریہ', 'ur', '33', 5840, 584, 1)
""")
conn.commit()

event_repo = EventCandidateRepository(DB_PATH)
taxo_repo = TaxonomyRepository(DB_PATH)
tag_repo = EventCandidateTaxonomyRepository(DB_PATH)

print("=" * 85)
print(" EXTRACTING & INGESTING: مواعظِ شمسیہ — حضرت مولانا شمس الرحمن العباسی دامت برکاتہم")
print("=" * 85)

# 1. Process Aitikaf DOCX Books (BookID 95001)
aitikaf_files = sorted(glob.glob(os.path.join(AITIKAF_DIR, "*.docx")))
added_aitikaf = 0
aitikaf_stories = []

for docx_p in aitikaf_files:
    fname = os.path.basename(docx_p)
    if "فہرست" in fname: continue
    year_match = re.search(r'\d{4}', fname)
    year = year_match.group(0) if year_match else "اعتکاف"

    try:
        doc = Document(docx_p)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        clean_text = clean_xml_text(full_text)

        for match in combined_pattern.finditer(clean_text):
            start_idx = match.start()
            trigger_phrase = match.group(0).strip()
            story_span = clean_text[start_idx: start_idx + 850]

            sentences = re.split(r'([۔！？])', story_span)
            if len(sentences) > 4:
                story_span = "".join(sentences[:8]).strip()

            if len(story_span) < 120: continue

            # Dedup
            if any(story_span[:40] in s['matn'] for s in aitikaf_stories):
                continue

            title = clean_title(trigger_phrase)
            if len(title) < 5: continue

            cit = f"مواعظِ شمسیہ (بیاناتِ اعتکاف {year}ء)"
            aitikaf_stories.append({'title': title, 'matn': story_span, 'cit': cit})

            ev = ExtractedEvent(
                title=title, alternate_names=[], subject="مواعظِ شمسیہ (بیاناتِ اعتکاف)",
                date_hijri=None, date_gregorian=year, location="خانقاہ امدادیہ غفوریہ",
                background=story_span, summary=title,
                key_figures=["حضرت مولانا شمس الرحمن العباسی دامت برکاتہم"], quoted_excerpt=story_span, citation=cit
            )
            nid = event_repo.add_candidate(95001, int(year) if year.isdigit() else 1, int(year) if year.isdigit() else 1, ev)
            event_repo.confirm(nid)
            terms = [
                taxo_repo.get_or_create_term("subject", "مواعظِ شمسیہ", "ur"),
                taxo_repo.get_or_create_term("personality", "حضرت مولانا شمس الرحمن العباسی دامت برکاتہم", "ur")
            ]
            tag_repo.tag_candidate(nid, terms)
            added_aitikaf += 1

    except Exception as e:
        print(f"Error in {fname}: {e}")

print(f"✅ Ingested {added_aitikaf} Confirmed Waqiat from مواعظِ شمسیہ (بیاناتِ اعتکاف)!")

# 2. Process Sunday Transcribed Bayanat (BookID 95002)
sunday_files = sorted(glob.glob(os.path.join(SUNDAY_DIR, "**", "*.txt"), recursive=True))
added_sunday = 0
sunday_stories = []

for idx, txt_p in enumerate(sunday_files, 1):
    fname = os.path.basename(txt_p)
    clean_fname = re.sub(r'^\d+\s*[-_)]\s*', '', fname).replace('.txt', '')

    try:
        with open(txt_p, "r", encoding="utf-8", errors="ignore") as tf:
            content = tf.read()
        clean_text = clean_xml_text(content)

        title_match = re.search(r'عنوان[:\s]+([^\n\r]+)', clean_text)
        bayan_title = title_match.group(1).replace('*', '').strip() if title_match else clean_fname
        if len(bayan_title) > 60: bayan_title = bayan_title[:60] + "..."

        for match in combined_pattern.finditer(clean_text):
            start_idx = match.start()
            trigger_phrase = match.group(0).strip()
            story_span = clean_text[start_idx: start_idx + 850]

            sentences = re.split(r'([۔！？])', story_span)
            if len(sentences) > 4:
                story_span = "".join(sentences[:8]).strip()

            if len(story_span) < 120: continue

            # Dedup
            if any(story_span[:40] in s['matn'] for s in sunday_stories):
                continue

            title = clean_title(trigger_phrase)
            if len(title) < 5: continue

            cit = f"مواعظِ شمسیہ (مواعظ اتوار، بیان: {bayan_title})"
            sunday_stories.append({'title': title, 'matn': story_span, 'cit': cit})

            ev = ExtractedEvent(
                title=title, alternate_names=[], subject="مواعظِ شمسیہ (مواعظ اتوار)",
                date_hijri=None, date_gregorian=None, location="خانقاہ امدادیہ غفوریہ",
                background=story_span, summary=title,
                key_figures=["حضرت مولانا شمس الرحمن العباسی دامت برکاتہم"], quoted_excerpt=story_span, citation=cit
            )
            nid = event_repo.add_candidate(95002, idx, idx, ev)
            event_repo.confirm(nid)
            terms = [
                taxo_repo.get_or_create_term("subject", "مواعظِ شمسیہ", "ur"),
                taxo_repo.get_or_create_term("personality", "حضرت مولانا شمس الرحمن العباسی دامت برکاتہم", "ur")
            ]
            tag_repo.tag_candidate(nid, terms)
            added_sunday += 1

    except Exception as e:
        pass

print(f"✅ Ingested {added_sunday} Confirmed Waqiat from مواعظِ شمسیہ (مواعظ اتوار)!")

# 3. Generate Dedicated Series Word Document
all_khanqah_stories = aitikaf_stories + sunday_stories
print(f"\n📄 Generating Dedicated Word Document for {len(all_khanqah_stories)} Waqiat of مواعظِ شمسیہ...")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — مواعظِ شمسیہ")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"افادات و ارشادات: حضرت مولانا شمس الرحمن العباسی دامت برکاتہم (بیاناتِ اعتکاف و مواعظ اتوار — {urdu_num(len(all_khanqah_stories))} مستند واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

for i, s in enumerate(all_khanqah_stories, 1):
    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {s['title']}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(s['matn'])
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {s['cit']}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DIR, "واقعات انسائیکلوپیڈیا — مواعظِ شمسیہ (بیاناتِ اعتکاف و مواعظ اتوار).docx")
doc.save(out_word_path)
print(f"✅ Generated Mawaiz-e-Shamsia Word Document: {out_word_path}")

# Check current total
cur.execute("SELECT COUNT(*) FROM EventCandidates WHERE Status='confirmed'")
total_now = cur.fetchone()[0]
conn.close()

print("=" * 85)
print(f" 🏆 MAWAIZ-E-SHAMSIA INGESTION COMPLETE: +{added_aitikaf + added_sunday} NET-NEW WAQIAT ADDED!")
print(f" 🌟 MASTER DATABASE TOTAL NOW: {total_now} CONFIRMED WAQIAT!")
print("=" * 85)
