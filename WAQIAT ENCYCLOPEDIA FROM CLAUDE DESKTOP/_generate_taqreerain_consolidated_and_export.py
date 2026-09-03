import os, sys, sqlite3, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r"F:\ISLAMIC RESEARCH HUB AI\src")

from islamic_research_hub.interfaces.catalog_export_cli import export_catalog_to_file
from islamic_research_hub.interfaces.book_package_export_cli import export_book_package_to_file

DB_PATH = Path(r"F:\ISLAMIC RESEARCH HUB AI\data\books.db")
OUTPUT_DOCX_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
OCR_DIR = r"F:\کتب\ocr text books\اصلاحی تقریریں"
PAGES_DIR = os.path.join(OCR_DIR, "pages")
MOBILE_ASSETS_DIR = Path(r"F:\ISLAMIC RESEARCH HUB AI\mobile\app\src\main\assets")
MOBILE_CATALOG_PATH = MOBILE_ASSETS_DIR / "catalog.db"
SAMPLE_BOOKS_DIR = MOBILE_ASSETS_DIR / "sample_books"

os.makedirs(OUTPUT_DOCX_DIR, exist_ok=True)
os.makedirs(OCR_DIR, exist_ok=True)

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

VOLUMES_MAP = {
    2: (4362, 2, 280, "اصلاحی تقریریں جلد 2"),
    3: (4473, 281, 521, "اصلاحی تقریریں جلد 3"),
    4: (4583, 522, 786, "اصلاحی تقریریں جلد 4"),
    5: (4594, 787, 1043, "اصلاحی تقریریں جلد 5"),
    6: (4604, 1044, 1282, "اصلاحی تقریریں جلد 6"),
    7: (4629, 1283, 1511, "اصلاحی تقریریں جلد 7"),
    8: (4680, 1512, 1728, "اصلاحی تقریریں جلد 8"),
    9: (4761, 1729, 1969, "اصلاحی تقریریں جلد 9")
}

# 1. Generate Consolidated Full Book and Volume TXT Files
print("=" * 85)
print(" 1. GENERATING CONSOLIDATED FULL BOOK & VOLUME TXT FILES")
print("=" * 85)
master_txt_path = os.path.join(OCR_DIR, "اصلاحی تقریریں (مکمل کتاب).txt")

with open(master_txt_path, "w", encoding="utf-8") as master_f:
    master_f.write("=== اصلاحی تقریریں (مکمل ۸ جلدیں — جلد ۲ تا ۹) ===\n")
    master_f.write("افادات: مفتی اعظم پاکستان حضرت مولانا مفتی محمد رفیع عثمانیؒ\n")
    master_f.write("مکمل متون برآمد شدہ از گوگل کلاؤڈ وژن OCR (۳۰۰ DPI)\n\n")

    for vol_num in range(2, 10):
        bid, sp, ep, vname = VOLUMES_MAP[vol_num]
        vol_txt_path = os.path.join(OCR_DIR, f"اصلاحی تقریریں - جلد {vol_num}.txt")
        with open(vol_txt_path, "w", encoding="utf-8") as vol_f:
            vol_f.write(f"═══ {vname} (صفحات ۱ تا {ep - sp + 1}) ═══\n\n")
            master_f.write(f"\n\n{'='*65}\n═══ {vname} ═══\n{'='*65}\n\n")

            for p in range(sp, ep + 1):
                pfile = os.path.join(PAGES_DIR, f"page_{p:04d}.txt")
                content = ""
                if os.path.exists(pfile):
                    with open(pfile, "r", encoding="utf-8") as pf:
                        content = pf.read().strip()

                vol_pno = p - sp + 1
                page_header = f"\n\n--- [جلد {vol_num}، صفحہ نمبر: {vol_pno}] ---\n"
                vol_f.write(page_header + content)
                master_f.write(page_header + content)

        print(f"  ✅ Saved: {os.path.basename(vol_txt_path)}")

master_size_mb = os.path.getsize(master_txt_path) / (1024 * 1024)
print(f"  🏆 Saved Master File: {os.path.basename(master_txt_path)} ({master_size_mb:.2f} MB)")

# 2. Generate Dedicated Word Document
print("\n" + "=" * 85)
print(" 2. GENERATING DEDICATED WORD DOCUMENT FOR ISLAHI TAQREERAIN")
print("=" * 85)

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

bids = [4362, 4473, 4583, 4594, 4604, 4629, 4680, 4761]
bids_str = ",".join(str(b) for b in bids)

cur.execute(f"""
    SELECT ec.EventCandidateID, ec.BookID, b.Title, ec.ChunkStartPage, ec.Title, ec.ExtractedDataJson
    FROM EventCandidates ec
    JOIN Books b ON ec.BookID = b.BookID
    WHERE ec.BookID IN ({bids_str}) AND ec.Status='confirmed'
    ORDER BY ec.BookID, ec.ChunkStartPage
""")
rows = cur.fetchall()

# Enrich books metadata in Books table
for v in range(2, 10):
    bid = VOLUMES_MAP[v][0]
    cur.execute("""
        UPDATE Books
        SET LibraryID=2, Publisher='ادارہ المعارف، کراچی',
            Category='خطبات و مواعظ', SeriesID=4362, VolumeNumber=?
        WHERE BookID=?
    """, (v, bid))
conn.commit()
conn.close()

print(f"Total Confirmed Waqiat for Islahi Taqreerain (8 Volumes): {len(rows)}")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — اصلاحی تقریریں")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"افادات: مفتی اعظم پاکستان حضرت مولانا مفتی محمد رفیع عثمانیؒ (مکمل ۸ جلدیں — {urdu_num(len(rows))} مستند و خالص واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

current_book_id = None

for i, r in enumerate(rows, 1):
    ev_id, bid, btitle, pno, title, djson = r
    data = json.loads(djson)
    matn = data.get('quoted_excerpt') or data.get('background') or ''
    cit = data.get('citation') or f"{btitle}، صفحہ {urdu_num(pno)}"

    if bid != current_book_id:
        current_book_id = bid
        p_vol = doc.add_paragraph()
        p_vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_vol = p_vol.add_run(f"═══ {btitle} ═══")
        r_vol.font.name = 'Jameel Noori Nastaleeq'
        r_vol.font.size = Pt(18)
        r_vol.font.bold = True
        r_vol.font.color.rgb = RGBColor(180, 50, 20)

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {title}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(matn)
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {cit}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DOCX_DIR, "واقعات انسائیکلوپیڈیا - اصلاحی تقریریں (مکمل ۸ جلدیں).docx")
doc.save(out_word_path)
print(f"✅ Successfully compiled: {out_word_path}")

# 3. Export Catalog and Book Packages
print("\n" + "=" * 85)
print(" 3. EXPORTING UPDATED CATALOG.DB AND BOOK PACKAGE FOR MOBILE")
print("=" * 85)

export_catalog_to_file(DB_PATH, MOBILE_CATALOG_PATH)
export_catalog_to_file(DB_PATH, Path("data/exports/catalog.db"))
print(f"✅ Exported mobile catalog: {MOBILE_CATALOG_PATH}")

out_pkg = SAMPLE_BOOKS_DIR / "book_4362.db"
export_book_package_to_file(DB_PATH, 4362, out_pkg)
pkg_size_mb = out_pkg.stat().st_size / (1024 * 1024)
print(f"✅ Exported Offline Reader Package: book_4362.db (اصلاحی تقریریں جلد 2) -> {pkg_size_mb:.2f} MB")
print("=" * 85)
