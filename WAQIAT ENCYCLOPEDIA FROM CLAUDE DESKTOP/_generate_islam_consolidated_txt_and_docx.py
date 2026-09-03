import os, sys, sqlite3, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DOCX_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
OCR_DIR = r"F:\کتب\ocr text books\اسلام اور ہماری زندگی"
PAGES_DIR = os.path.join(OCR_DIR, "pages")

os.makedirs(OUTPUT_DOCX_DIR, exist_ok=True)
os.makedirs(OCR_DIR, exist_ok=True)

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

VOLUMES_MAP = {
    1: (3392, 1, 345, "اسلام اور ہماری زندگی جلد 1"),
    2: (3465, 346, 682, "اسلام اور ہماری زندگی جلد 2"),
    3: (3523, 683, 1051, "اسلام اور ہماری زندگی جلد 3"),
    4: (3633, 1052, 1356, "اسلام اور ہماری زندگی جلد 4"),
    5: (3744, 1357, 1701, "اسلام اور ہماری زندگی جلد 5"),
    6: (3854, 1702, 2014, "اسلام اور ہماری زندگی جلد 6"),
    7: (3961, 2015, 2367, "اسلام اور ہماری زندگی جلد 7"),
    8: (4051, 2368, 2720, "اسلام اور ہماری زندگی جلد 8"),
    9: (4146, 2721, 3017, "اسلام اور ہماری زندگی جلد 9"),
    10: (4251, 3018, 3306, "اسلام اور ہماری زندگی جلد 10")
}

# 1. Generate Consolidated Full Book and Volume TXT Files
print("=" * 85)
print(" 1. GENERATING CONSOLIDATED FULL BOOK & VOLUME TXT FILES")
print("=" * 85)
master_txt_path = os.path.join(OCR_DIR, "اسلام اور ہماری زندگی (مکمل کتاب).txt")

with open(master_txt_path, "w", encoding="utf-8") as master_f:
    master_f.write("=== اسلام اور ہماری زندگی (مکمل ۱۰ جلدیں) ===\n")
    master_f.write("افادات: شیخ الاسلام مفتی محمد تقی عثمانی مدظلہم\n")
    master_f.write("مکمل متون برآمد شدہ از گوگل کلاؤڈ وژن OCR (۳۰۰ DPI)\n\n")

    for vol_num in range(1, 11):
        bid, sp, ep, vname = VOLUMES_MAP[vol_num]
        vol_txt_path = os.path.join(OCR_DIR, f"اسلام اور ہماری زندگی - جلد {vol_num}.txt")
        with open(vol_txt_path, "w", encoding="utf-8") as vol_f:
            vol_f.write(f"═══ {vname} (صفحات ۱ تا {ep - sp + 1}) ═══\n\n")
            master_f.write(f"\n\n{'='*65}\n═══ {vname} ═══\n{'='*65}\n\n")

            for p in range(sp, ep + 1):
                pfile = os.path.join(PAGES_DIR, f"page_{p:04d}.txt")
                content = ""
                if os.path.exists(pfile):
                    with open(pfile, "r", encoding="utf-8") as pf:
                        content = pf.read().strip()

                vol_pno = p - sp + 1
                page_header = f"\n\n--- [جلد {vol_num}، صفحہ نمبر: {vol_pno}] ---\n"
                vol_f.write(page_header + content)
                master_f.write(page_header + content)

        print(f"  ✅ Saved: {os.path.basename(vol_txt_path)}")

master_size_mb = os.path.getsize(master_txt_path) / (1024 * 1024)
print(f"  🏆 Saved Master File: {os.path.basename(master_txt_path)} ({master_size_mb:.2f} MB)")

# 2. Generate Dedicated Word Document
print("\n" + "=" * 85)
print(" 2. GENERATING DEDICATED WORD DOCUMENT FOR ISLAM AUR HAMARI ZINDAGI")
print("=" * 85)

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

bids = [3392, 3465, 3523, 3633, 3744, 3854, 3961, 4051, 4146, 4251]
bids_str = ",".join(str(b) for b in bids)

cur.execute(f"""
    SELECT ec.EventCandidateID, ec.BookID, b.Title, ec.ChunkStartPage, ec.Title, ec.ExtractedDataJson
    FROM EventCandidates ec
    JOIN Books b ON ec.BookID = b.BookID
    WHERE ec.BookID IN ({bids_str}) AND ec.Status='confirmed'
    ORDER BY ec.BookID, ec.ChunkStartPage
""")
rows = cur.fetchall()
conn.close()

print(f"Total Confirmed Waqiat for Islam aur Hamari Zindagi (10 Volumes): {len(rows)}")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — اسلام اور ہماری زندگی")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"افادات: شیخ الاسلام حضرت مولانا مفتی محمد تقی عثمانی مدظلہم (مکمل ۱۰ جلدیں — {urdu_num(len(rows))} مستند و خالص واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

current_book_id = None

for i, r in enumerate(rows, 1):
    ev_id, bid, btitle, pno, title, djson = r
    data = json.loads(djson)
    matn = data.get('quoted_excerpt') or data.get('background') or ''
    cit = data.get('citation') or f"{btitle}، صفحہ {urdu_num(pno)}"

    if bid != current_book_id:
        current_book_id = bid
        p_vol = doc.add_paragraph()
        p_vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_vol = p_vol.add_run(f"═══ {btitle} ═══")
        r_vol.font.name = 'Jameel Noori Nastaleeq'
        r_vol.font.size = Pt(18)
        r_vol.font.bold = True
        r_vol.font.color.rgb = RGBColor(180, 50, 20)

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {title}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(matn)
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {cit}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DOCX_DIR, "واقعات انسائیکلوپیڈیا - اسلام اور ہماری زندگی (مکمل ۱۰ جلدیں).docx")
doc.save(out_word_path)
print(f"✅ Successfully compiled: {out_word_path}")
print("=" * 85)
