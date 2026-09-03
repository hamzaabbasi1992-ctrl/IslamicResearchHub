import os, sys, glob
from docx import Document
sys.stdout.reconfigure(encoding='utf-8')

aitikaf_dir = r"E:\KHANQAH\AITIKAAF BAYANAT\book formatted"
sunday_dir = r"E:\KHANQAH\SUNDAY BAYANS (FINAL)\bayan to text gemini output"

print("=" * 85)
print(" INSPECTING AITIKAF BAYANAT DOCX FILES")
print("=" * 85)
aitikaf_files = sorted(glob.glob(os.path.join(aitikaf_dir, "*.docx")))
print(f"Total Aitikaf DOCX Books: {len(aitikaf_files)}")
for f in aitikaf_files:
    fname = os.path.basename(f)
    size_mb = os.path.getsize(f) / (1024*1024)
    try:
        doc = Document(f)
        total_p = len(doc.paragraphs)
        sample = next((p.text[:80] for p in doc.paragraphs if len(p.text.strip()) > 30), "")
        print(f"  📖 {fname:45s} | {size_mb:5.2f} MB | {total_p:4d} paras | Sample: {sample}...")
    except Exception as e:
        print(f"  ❌ {fname:45s} | Error: {e}")

print("\n" + "=" * 85)
print(" INSPECTING SUNDAY BAYANAT TXT FILES")
print("=" * 85)
sunday_files = sorted(glob.glob(os.path.join(sunday_dir, "**", "*.txt"), recursive=True))
print(f"Total Sunday Transcribed Bayanat Files: {len(sunday_files)}")
total_chars = 0
for f in sunday_files[:15]:
    fname = os.path.basename(f)
    size_kb = os.path.getsize(f) / 1024
    with open(f, "r", encoding="utf-8", errors="ignore") as tf:
        text = tf.read().strip()
        total_chars += len(text)
        sample = text[:80].replace("\n", " ")
        print(f"  📝 {fname:55s} | {size_kb:6.1f} KB | Sample: {sample}...")

if len(sunday_files) > 15:
    print(f"  ... and {len(sunday_files)-15} more Sunday Bayanat text files")
print("=" * 85)
