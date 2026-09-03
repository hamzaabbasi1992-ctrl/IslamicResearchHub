import sys, os, sqlite3, json, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
os.makedirs(OUTPUT_DIR, exist_ok=True)

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

def clean_xml_text(s):
    if not s: return ""
    s = re.sub(r'</?[^>]+>', ' ', s)
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', s)
    s = re.sub(r'www\.[a-zA-Z0-9_\-\.]+', '', s)
    return re.sub(r'\s+', ' ', s).strip()

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

# Get all confirmed Waqiat in database
cur.execute("""
    SELECT ec.EventCandidateID, ec.BookID, b.Title, ec.ChunkStartPage, ec.ExtractedDataJson
    FROM EventCandidates ec
    JOIN Books b ON ec.BookID = b.BookID
    WHERE ec.Status = 'confirmed'
    ORDER BY b.Title, ec.ChunkStartPage
""")
rows = cur.fetchall()
conn.close()

total_count = len(rows)
print(f"Total Confirmed Waqiat in DB for Grand Master: {total_count}")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Main Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — گرینڈ ماسٹر انسائیکلوپیڈیا")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"جامع انسائیکلوپیڈیا برائے اسلامی اصلاحی، تاریخی و تفسیری واقعات (تمام کتب و سیریز — {urdu_num(total_count)} مستند واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

for i, (cid, bid, btitle, pno, data_json) in enumerate(rows, 1):
    data = json.loads(data_json) if data_json else {}
    w_title = clean_xml_text(data.get('title') or f"واقعہ {urdu_num(i)}")
    w_matn = clean_xml_text(data.get('quoted_excerpt') or data.get('background') or '')
    w_cit = data.get('citation') or f"{btitle}، صفحہ {urdu_num(pno)}"

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {w_title}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(w_matn)
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {w_cit}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_file = os.path.join(OUTPUT_DIR, f"واقعات انسائیکلوپیڈیا — گرینڈ ماسٹر انسائیکلوپیڈیا (تمام کتب و سیریز — {total_count} واقعات).docx")
doc.save(out_file)
print(f"✅ Generated Grand Master Word Document: {out_file} ({total_count} Waqiat)")
