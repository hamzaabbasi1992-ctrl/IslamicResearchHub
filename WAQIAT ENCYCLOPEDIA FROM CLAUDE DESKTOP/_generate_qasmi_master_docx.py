import os, sys, sqlite3, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
os.makedirs(OUTPUT_DIR, exist_ok=True)

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

cur.execute("""
    SELECT ec.EventCandidateID, ec.BookID, b.Title, ec.ChunkStartPage, ec.Title, ec.ExtractedDataJson
    FROM EventCandidates ec
    JOIN Books b ON ec.BookID = b.BookID
    WHERE ec.BookID IN (3534, 3545, 35451, 35452, 3556, 3567) AND ec.Status='confirmed'
    ORDER BY ec.BookID, ec.ChunkStartPage
""")
rows = cur.fetchall()
conn.close()

print(f"Total Confirmed Waqiat for Khutbat Qasmi (6 Volumes): {len(rows)}")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — خطباتِ قاسمی")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"افادات: خطیبِ اسلام حضرت مولانا ضیاء القاسمؒ (مکمل ۶ جلدیں — {urdu_num(len(rows))} مستند و خالص واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

current_book_id = None
vol_counter = 0

for i, r in enumerate(rows, 1):
    ev_id, bid, btitle, pno, title, djson = r
    data = json.loads(djson)
    matn = data.get('quoted_excerpt') or data.get('background') or ''
    cit = data.get('citation') or f"{btitle}، صفحہ {urdu_num(pno)}"

    if bid != current_book_id:
        current_book_id = bid
        vol_counter += 1
        p_vol = doc.add_paragraph()
        p_vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_vol = p_vol.add_run(f"═══ {btitle} ═══")
        r_vol.font.name = 'Jameel Noori Nastaleeq'
        r_vol.font.size = Pt(18)
        r_vol.font.bold = True
        r_vol.font.color.rgb = RGBColor(180, 50, 20)

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {title}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(matn)
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {cit}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DIR, "واقعات انسائیکلوپیڈیا - خطبات قاسمی (مکمل ۶ جلدیں).docx")
doc.save(out_word_path)
print(f"✅ Successfully compiled: {out_word_path}")
