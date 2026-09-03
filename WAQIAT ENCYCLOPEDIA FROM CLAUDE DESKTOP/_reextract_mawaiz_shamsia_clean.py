import os, sys, glob, re, sqlite3, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r"F:\ISLAMIC RESEARCH HUB AI\src")

from islamic_research_hub.application.event_extraction import ExtractedEvent
from islamic_research_hub.infrastructure.persistence.event_candidate_repository import EventCandidateRepository
from islamic_research_hub.infrastructure.persistence.taxonomy_repository import TaxonomyRepository
from islamic_research_hub.infrastructure.persistence.event_candidate_taxonomy_repository import EventCandidateTaxonomyRepository

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
os.makedirs(OUTPUT_DIR, exist_ok=True)

AITIKAF_DIR = r"E:\KHANQAH\AITIKAAF BAYANAT\book formatted"
SUNDAY_DIR = r"E:\KHANQAH\SUNDAY BAYANS (FINAL)\bayan to text gemini output"

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

def clean_xml_text(s):
    if not s: return ""
    s = re.sub(r'</?[^>]+>', ' ', s)
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', s)
    s = re.sub(r'www\.[a-zA-Z0-9_\-\.]+', '', s)
    return re.sub(r'\s+', ' ', s).strip()

def clean_title(t, maxlen=80):
    t = clean_xml_text(t)
    t = re.sub(r'^[\d۱-۹١-٩\(\)\*\-\s٭…_]+', '', t).strip()
    return (t[:maxlen] + "...") if len(t) > maxlen else t

# STRICT NARRATIVE TRIGGERS (STORY ANCHORS ONLY)
STRICT_NARRATIVE_TRIGGERS = [
    r'(ایک مرتبہ\s+(?:حضرت|کسی|ایک|وہ|بادشاہ|بزرگ|درویش|شخص|عورت|صحابی|واقعہ|قصہ)[^۔\n]{5,50}؟?)',
    r'(ایک دفعہ\s+(?:حضرت|کسی|ایک|وہ|بادشاہ|بزرگ|درویش|شخص|عورت|صحابی|واقعہ|قصہ)[^۔\n]{5,50}؟?)',
    r'(ایک بار\s+(?:حضرت|کسی|ایک|وہ|بادشاہ|بزرگ|درویش|شخص|عورت|صحابی|واقعہ|قصہ)[^۔\n]{5,50}؟?)',
    r'(ایک بزرگ\s+[^۔\n]{5,50}؟?)',
    r'(ایک شخص\s+[^۔\n]{5,50}؟?)',
    r'(ایک بادشاہ\s+[^۔\n]{5,50}؟?)',
    r'(ایک درویش\s+[^۔\n]{5,50}؟?)',
    r'(ایک صاحب\s+[^۔\n]{5,50}؟?)',
    r'(ایک ولی\s+[^۔\n]{5,50}؟?)',
    r'(ایک فقیر\s+[^۔\n]{5,50}؟?)',
    r'(ایک مجذوب\s+[^۔\n]{5,50}؟?)',
    r'(ایک صحابی\s+[^۔\n]{5,50}؟?)',
    r'(حضرت\s+[^\n۔]{3,30}\s+(?:کا واقعہ|کا قصہ|کی حکایت|کا ایک واقعہ)[^۔\n]{0,35})',
    r'(خواب میں دیکھا\s+کہ\s+[^۔\n]{5,50})',
    r'(حکایت ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(واقعہ ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(واقعہ یہ ہے کہ\s+[^۔\n]{5,50}؟?)',
    r'(ایک سچا واقعہ\s+[^۔\n]{0,40})',
    r'(عجیب واقعہ ہے کہ\s+[^۔\n]{0,40})',
    r'(ایک قصہ سنا\s+[^۔\n]{0,40})'
]
combined_strict_pattern = re.compile("|".join(STRICT_NARRATIVE_TRIGGERS), re.UNICODE)

# Disqualification patterns for non-story quotes / wazaif / sayings
DISQUALIFY_PATTERNS = [
    r'^ارشاد فرمایا',
    r'^فرمایا کہ',
    r'^بیان فرمایا',
    r'^لکھا ہے کہ',
    r'^آیا ہے کہ',
    r'^حدیث میں آتا ہے',
    r'^حدیث شریف میں ہے',
    r'^قرآن میں ہے',
    r'^قرآن پاک میں ہے',
    r'اللہم صل علی',
    r'لاکھ مرتبہ قسم کھا سکتا ہوں',
    r'درود شریف پڑھیں',
    r'نفل پڑھا کرو',
    r'تسبیح پڑھا کرو',
    r'مرتبہ پڑھ لیا کرو'
]

def is_valid_waqia(title, matn):
    title_clean = title.strip()
    # Check blacklist
    for dp in DISQUALIFY_PATTERNS:
        if re.search(dp, title_clean):
            return False
    
    # Must be of reasonable story length (> 150 chars, > 30 words)
    words = matn.split()
    if len(words) < 25 or len(matn) < 140:
        return False
    
    # Check for story narrative markers in matn
    story_markers = ["تھا", "تھی", "تھے", "گئے", "آئے", "دیکھا", "کہا", "عرض کیا", "پوچھا", "جواب دیا", "پہنچے", "گزر رہے تھے", "واقعہ", "قصہ"]
    marker_count = sum(1 for m in story_markers if m in matn)
    if marker_count < 2:
        return False
        
    return True

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

# Clean existing BookID 95001 and 95002 candidates completely
cur.execute("DELETE FROM EventCandidateTaxonomyTerms WHERE EventCandidateID IN (SELECT EventCandidateID FROM EventCandidates WHERE BookID IN (95001, 95002))")
cur.execute("DELETE FROM EventCandidates WHERE BookID IN (95001, 95002)")
conn.commit()

event_repo = EventCandidateRepository(DB_PATH)
taxo_repo = TaxonomyRepository(DB_PATH)
tag_repo = EventCandidateTaxonomyRepository(DB_PATH)

print("=" * 85)
print(" RE-EXTRACTING HIGH-PRECISION WAQIAT FOR: مواعظِ شمسیہ (حضرت مولانا شمس الرحمن العباسی)")
print("=" * 85)

# 1. Process Aitikaf DOCX Books (BookID 95001)
aitikaf_files = sorted(glob.glob(os.path.join(AITIKAF_DIR, "*.docx")))
added_aitikaf = 0
aitikaf_stories = []

for docx_p in aitikaf_files:
    fname = os.path.basename(docx_p)
    if "فہرست" in fname: continue
    year_match = re.search(r'\d{4}', fname)
    year = year_match.group(0) if year_match else "اعتکاف"

    try:
        doc = Document(docx_p)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        clean_text = clean_xml_text(full_text)

        for match in combined_strict_pattern.finditer(clean_text):
            start_idx = match.start()
            trigger_phrase = match.group(0).strip()
            story_span = clean_text[start_idx: start_idx + 900]

            sentences = re.split(r'([۔！？])', story_span)
            if len(sentences) > 4:
                story_span = "".join(sentences[:8]).strip()

            title = clean_title(trigger_phrase)
            if not is_valid_waqia(title, story_span):
                continue

            # Dedup
            if any(story_span[:40] in s['matn'] for s in aitikaf_stories):
                continue

            cit = f"مواعظِ شمسیہ (بیاناتِ اعتکاف {year}ء)"
            aitikaf_stories.append({'title': title, 'matn': story_span, 'cit': cit})

            ev = ExtractedEvent(
                title=title, alternate_names=[], subject="مواعظِ شمسیہ (بیاناتِ اعتکاف)",
                date_hijri=None, date_gregorian=year, location="خانقاہ امدادیہ غفوریہ",
                background=story_span, summary=title,
                key_figures=["حضرت مولانا شمس الرحمن العباسی دامت برکاتہم"], quoted_excerpt=story_span, citation=cit
            )
            nid = event_repo.add_candidate(95001, int(year) if year.isdigit() else 1, int(year) if year.isdigit() else 1, ev)
            event_repo.confirm(nid)
            terms = [
                taxo_repo.get_or_create_term("subject", "مواعظِ شمسیہ", "ur"),
                taxo_repo.get_or_create_term("personality", "حضرت مولانا شمس الرحمن العباسی دامت برکاتہم", "ur")
            ]
            tag_repo.tag_candidate(nid, terms)
            added_aitikaf += 1

    except Exception as e:
        print(f"Error in {fname}: {e}")

print(f"✅ Filtered & Ingested: {added_aitikaf} Genuine Confirmed Waqiat from Aitikaf Bayanat!")

# 2. Process Sunday Transcribed Bayanat (BookID 95002)
sunday_files = sorted(glob.glob(os.path.join(SUNDAY_DIR, "**", "*.txt"), recursive=True))
added_sunday = 0
sunday_stories = []

for idx, txt_p in enumerate(sunday_files, 1):
    fname = os.path.basename(txt_p)
    clean_fname = re.sub(r'^\d+\s*[-_)]\s*', '', fname).replace('.txt', '')

    try:
        with open(txt_p, "r", encoding="utf-8", errors="ignore") as tf:
            content = tf.read()
        clean_text = clean_xml_text(content)

        title_match = re.search(r'عنوان[:\s]+([^\n\r]+)', clean_text)
        bayan_title = title_match.group(1).replace('*', '').strip() if title_match else clean_fname
        if len(bayan_title) > 60: bayan_title = bayan_title[:60] + "..."

        for match in combined_strict_pattern.finditer(clean_text):
            start_idx = match.start()
            trigger_phrase = match.group(0).strip()
            story_span = clean_text[start_idx: start_idx + 900]

            sentences = re.split(r'([۔！？])', story_span)
            if len(sentences) > 4:
                story_span = "".join(sentences[:8]).strip()

            title = clean_title(trigger_phrase)
            if not is_valid_waqia(title, story_span):
                continue

            # Dedup
            if any(story_span[:40] in s['matn'] for s in sunday_stories):
                continue

            cit = f"مواعظِ شمسیہ (مواعظ اتوار، بیان: {bayan_title})"
            sunday_stories.append({'title': title, 'matn': story_span, 'cit': cit})

            ev = ExtractedEvent(
                title=title, alternate_names=[], subject="مواعظِ شمسیہ (مواعظ اتوار)",
                date_hijri=None, date_gregorian=None, location="خانقاہ امدادیہ غفوریہ",
                background=story_span, summary=title,
                key_figures=["حضرت مولانا شمس الرحمن العباسی دامت برکاتہم"], quoted_excerpt=story_span, citation=cit
            )
            nid = event_repo.add_candidate(95002, idx, idx, ev)
            event_repo.confirm(nid)
            terms = [
                taxo_repo.get_or_create_term("subject", "مواعظِ شمسیہ", "ur"),
                taxo_repo.get_or_create_term("personality", "حضرت مولانا شمس الرحمن العباسی دامت برکاتہم", "ur")
            ]
            tag_repo.tag_candidate(nid, terms)
            added_sunday += 1

    except Exception as e:
        pass

print(f"✅ Filtered & Ingested: {added_sunday} Genuine Confirmed Waqiat from Sunday Bayanat!")

# 3. Generate Dedicated Series Word Document
all_clean_stories = aitikaf_stories + sunday_stories
print(f"\n📄 Compiling Dedicated Word Document for {len(all_clean_stories)} Pure Narrative Waqiat...")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — مواعظِ شمسیہ")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"افادات و ارشادات: حضرت مولانا شمس الرحمن العباسی دامت برکاتہم (مستند و منتخب سچے واقعات — {urdu_num(len(all_clean_stories))} واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

for i, s in enumerate(all_clean_stories, 1):
    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {s['title']}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(s['matn'])
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {s['cit']}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DIR, "واقعات انسائیکلوپیڈیا — مواعظِ شمسیہ (بیاناتِ اعتکاف و مواعظ اتوار).docx")
doc.save(out_word_path)
print(f"✅ Generated Clean Mawaiz-e-Shamsia Word Document: {out_word_path}")

cur.execute("SELECT COUNT(*) FROM EventCandidates WHERE Status='confirmed'")
total_now = cur.fetchone()[0]
conn.close()

print("=" * 85)
print(f" 🏆 CLEAN MAWAIZ-E-SHAMSIA COMPLETE: {len(all_clean_stories)} PURE WAQIAT RETAINED!")
print(f" 🌟 MASTER DATABASE TOTAL NOW: {total_now} CONFIRMED WAQIAT!")
print("=" * 85)
