import os, sys, sqlite3, json, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
sys.stdout.reconfigure(encoding='utf-8')

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
OUTPUT_DIR = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP\مکمل جلدیں واقعات کتابوں سے"
os.makedirs(OUTPUT_DIR, exist_ok=True)

URDU_DIGITS = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
def urdu_num(n): return str(n).translate(URDU_DIGITS)

def clean_xml_text(s):
    if not s: return ""
    s = re.sub(r'</?[^>]+>', ' ', s)
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', s)
    s = re.sub(r'www\.[a-zA-Z0-9_\-\.]+', '', s)
    return re.sub(r'\s+', ' ', s).strip()

conn = sqlite3.connect(DB_PATH)
cur = conn.cursor()

hakim_bids = [5091, 5102, 5113, 5123, 5128]
placeholders = ",".join(str(b) for b in hakim_bids)

cur.execute(f"""
    SELECT ec.EventCandidateID, ec.BookID, b.Title, ec.ChunkStartPage, ec.ExtractedDataJson
    FROM EventCandidates ec
    JOIN Books b ON ec.BookID = b.BookID
    WHERE ec.BookID IN ({placeholders}) AND ec.Status = 'confirmed'
    ORDER BY ec.BookID, ec.ChunkStartPage, ec.EventCandidateID
""")
rows = cur.fetchall()
conn.close()

print(f"Total Confirmed Waqiat for Khutbat Hakim ul Islam (5 Volumes): {len(rows)}")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("واقعات انسائیکلوپیڈیا — خطباتِ حکیم الاسلام")
r_title.font.name = 'Jameel Noori Nastaleeq'
r_title.font.size = Pt(26)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(16, 78, 139)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run(f"خطیبِ اسلام حضرت مولانا قاری محمد طیب قاسمیؒ (مکمل ۵ جلدیں — {urdu_num(len(rows))} مستند واقعات)")
r_sub.font.name = 'Jameel Noori Nastaleeq'
r_sub.font.size = Pt(15)
r_sub.font.color.rgb = RGBColor(100, 100, 100)

p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_div = p_div.add_run("❖ ❖ ❖")
r_div.font.size = Pt(14)
r_div.font.color.rgb = RGBColor(180, 140, 20)

current_book_id = None
vol_num = 0

for i, r in enumerate(rows, 1):
    ev_id, bid, btitle, pnum, data_json = r
    try:
        data = json.loads(data_json)
    except:
        continue
    
    title = clean_xml_text(data.get("title", f"واقعہ {i}"))
    matn = clean_xml_text(data.get("quoted_excerpt") or data.get("background") or "")
    cit = clean_xml_text(data.get("citation", f"{btitle}، ص {pnum}"))

    if bid != current_book_id:
        current_book_id = bid
        vol_num += 1
        p_vol = doc.add_paragraph()
        p_vol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_vol = p_vol.add_run(f"═══ {btitle} (جلد {urdu_num(vol_num)}) ═══")
        r_vol.font.name = 'Jameel Noori Nastaleeq'
        r_vol.font.size = Pt(18)
        r_vol.font.bold = True
        r_vol.font.color.rgb = RGBColor(139, 69, 19)

    # Heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run(f"✦ واقعہ نمبر {urdu_num(i)}: {title}")
    r_head.font.name = 'Jameel Noori Nastaleeq'
    r_head.font.size = Pt(14)
    r_head.font.bold = True
    r_head.font.color.rgb = RGBColor(20, 90, 50)

    # Body
    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(matn)
    r_body.font.name = 'Jameel Noori Nastaleeq'
    r_body.font.size = Pt(12)

    # Citation
    p_cit = doc.add_paragraph()
    p_cit.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_cit = p_cit.add_run(f"حوالہ: {cit}")
    r_cit.font.name = 'Jameel Noori Nastaleeq'
    r_cit.font.size = Pt(10)
    r_cit.font.italic = True
    r_cit.font.color.rgb = RGBColor(120, 120, 120)

    # Divider
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = p_sep.add_run("❖ ❖ ❖")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = RGBColor(200, 180, 140)

out_word_path = os.path.join(OUTPUT_DIR, "واقعات انسائیکلوپیڈیا - خطبات حکیم الاسلام (مکمل ۵ جلدیں).docx")
doc.save(out_word_path)
print(f"✅ Successfully compiled: {out_word_path}")
