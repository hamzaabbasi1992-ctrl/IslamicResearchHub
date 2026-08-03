"""Local .docx storage adapter for Research Notes, using python-docx.

The one concrete storage backend this feature ships with - implements
the `NotesStorage` shape `ResearchNotesManager` depends on structurally
(see `research_notes_manager.py`). A future cloud backend (Google Docs,
Google Drive, OneDrive, Dropbox) would be a second class with the same
three methods, swapped in wherever `LocalDocxStorage()` is constructed
today, with no change to the manager or the dialog.
"""

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from islamic_research_hub.research_notes.research_notes_manager import Quotation

NOTES_FOLDER_NAME = "Maktaba Research Notes"


class NoteFileLockedError(Exception):
    """Raised when a note document can't be saved because it's open
    elsewhere. Microsoft Word holds a real exclusive lock on an open
    .docx file on Windows, which turns a plain save into a
    `PermissionError` - this is that same failure, given a name callers
    can catch and show a friendly message for instead of crashing."""


class LocalDocxStorage:
    """Store every note document as a real .docx file under the user's
    Documents folder."""

    def notes_folder(self) -> Path:
        """Return the real notes folder, creating it if it doesn't exist yet."""
        folder = Path.home() / "Documents" / NOTES_FOLDER_NAME
        folder.mkdir(parents=True, exist_ok=True)
        return folder

    def list_documents(self) -> tuple[Path, ...]:
        """Return every real .docx note file in the notes folder, alphabetically."""
        return tuple(sorted(self.notes_folder().glob("*.docx")))

    def create_document(self, name: str) -> Path:
        """Create a new, empty .docx note file and return its real path.

        `name` is sanitized to a safe filename (a real thesis/book title
        can carry characters Windows filenames don't allow) and a numeric
        suffix is added on a real collision, so an existing document is
        never overwritten by creating a new one with the same name.
        """
        safe_name = _sanitize_filename(name) or "Untitled"
        path = _unique_path(self.notes_folder() / f"{safe_name}.docx")
        Document().save(path)
        return path

    def append_quotation(self, path: Path, quotation: Quotation) -> None:
        """Append one quotation entry to the end of an existing document.

        The real document is opened, a new entry is appended after
        whatever's already there, and the whole document is saved back -
        existing content is never overwritten. Raises
        `NoteFileLockedError` (never crashes) if the document is
        currently open in Word.
        """
        try:
            document = Document(path)
        except PackageNotFoundError:
            document = Document()

        for line in _quotation_lines(quotation):
            document.add_paragraph(line)

        try:
            document.save(path)
        except PermissionError as error:
            raise NoteFileLockedError(
                "This note file is currently open.\nPlease close it and try again."
            ) from error


def _quotation_lines(quotation: Quotation) -> tuple[str, ...]:
    """Build the append block's lines, matching the feature's fixed format
    (label on its own line, value below it, a blank line between fields)."""
    lines: list[str] = ["Book:", quotation.book_title, ""]
    lines += ["Author:", quotation.author or "Unknown", ""]
    if quotation.volume is not None:
        lines += ["Volume:", str(quotation.volume), ""]
    if quotation.chapter:
        lines += ["Chapter:", quotation.chapter, ""]
    if quotation.page_number is not None:
        lines += ["Page:", str(quotation.page_number), ""]
    lines += ["Date:", date.today().isoformat(), ""]
    lines += ["-" * 40, ""]
    lines += ["Selected Text", "", quotation.selected_text, ""]
    lines += ["-" * 40, ""]
    lines += ["My Notes", "", "_" * 36, ""]
    return tuple(lines)


def _sanitize_filename(name: str) -> str:
    """Strip characters Windows filenames don't allow, collapse whitespace."""
    cleaned = re.sub(r'[\\/:*?"<>|]', "", name).strip()
    return re.sub(r"\s+", " ", cleaned)


def _unique_path(path: Path) -> Path:
    """Return `path`, or `<name> (2).docx` / `(3).docx`... on a real name
    collision - an existing document is never silently overwritten."""
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    counter = 2
    while True:
        candidate = path.with_name(f"{stem} ({counter}){suffix}")
        if not candidate.exists():
            return candidate
        counter += 1
