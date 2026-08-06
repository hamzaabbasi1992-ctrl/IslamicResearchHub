"""Export a real named collection (Phase 14 Milestone 1) as a real .docx
document, one section per item with its real citation.

Lives alongside `docx_writer.py` (the other real python-docx-touching
code in this project) rather than under `application/` - both are real
storage-adjacent adapters, not pure domain logic.
"""

from dataclasses import dataclass
from pathlib import Path

from docx import Document

from islamic_research_hub.shared.citation_formatting import format_citation


@dataclass(frozen=True, slots=True)
class CollectionExportItem:
    """One real collection item, fully hydrated with what the export
    needs - the caller (a desktop screen) is responsible for fetching
    this from `CollectionRepository`/`BookBrowserRepository`, keeping
    this module itself free of any database access."""

    book_title: str
    author: str | None
    volume_number: int | None
    page_number: int
    content: str


def build_export_document(collection_name: str, items: tuple[CollectionExportItem, ...]) -> Document:
    """Return a real, in-memory .docx `Document` for one collection - a
    title page, then one real section per item (heading, citation, real
    page content), in the same order `items` was given.

    Returns the `Document` object rather than saving it, so this stays
    testable without touching a real file - `export_collection_to_docx()`
    below is the thin save wrapper real callers use.
    """
    document = Document()
    document.add_heading(collection_name, level=1)
    if not items:
        document.add_paragraph("This collection has no items yet.")
        return document

    for item in items:
        document.add_heading(item.book_title, level=2)
        citation = format_citation(
            item.book_title, item.page_number, paragraph_index=1, volume_number=item.volume_number
        )
        meta_bits = [item.author or "Unknown author", citation]
        document.add_paragraph(" - ".join(meta_bits))
        document.add_paragraph(item.content)
    return document


def export_collection_to_docx(
    collection_name: str, items: tuple[CollectionExportItem, ...], output_path: Path
) -> None:
    """Build and save a real collection export document to `output_path`."""
    build_export_document(collection_name, items).save(output_path)
