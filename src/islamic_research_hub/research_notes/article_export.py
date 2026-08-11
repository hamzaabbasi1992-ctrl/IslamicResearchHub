"""Export a real, already-composed article/book chapter as a real .docx document.

Deliberately does not gather any evidence or write any prose itself - it
formats sections and a works-cited list the caller already composed
(e.g. an AI assistant, working from real `search_text`/`get_citation`
results) into a real, shareable document. Lives alongside
`docx_writer.py`/`ai_answer_export.py`/`collection_export.py` (the other
real python-docx-touching code in this project), not under
`application/` - all are real storage-adjacent adapters, not pure domain
logic.
"""

from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True, slots=True)
class ArticleSection:
    """One real heading + body of an already-composed article/chapter."""

    heading: str
    body: str


def build_article_document(
    title: str, sections: tuple[ArticleSection, ...], sources: tuple[str, ...] = ()
) -> Document:
    """Return a real, in-memory .docx `Document` for one already-composed
    article - a title, one real section per (heading, body) pair in the
    given order, then a real "Works Cited" list of already-formatted
    citation strings (e.g. from `format_citation()`), if any were given.

    Returns the `Document` object rather than saving it, so this stays
    testable without touching a real file - `export_article_to_docx()`
    below is the thin save wrapper real callers use.
    """
    document = Document()
    document.add_heading(title, level=1)
    for section in sections:
        if section.heading:
            document.add_heading(section.heading, level=2)
        document.add_paragraph(section.body)
    if sources:
        document.add_heading("Works Cited", level=2)
        for source in sources:
            document.add_paragraph(f"- {source}")
    return document


def export_article_to_docx(
    title: str,
    sections: tuple[ArticleSection, ...],
    sources: tuple[str, ...],
    output_path: Path,
) -> None:
    """Build and save a real article export document to `output_path`."""
    build_article_document(title, sections, sources).save(output_path)
