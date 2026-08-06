"""Export a real AI Assistant answer (Phase 11) as a real .docx document
(Phase 16 Milestone 1: AI content generator).

Deliberately does not gather any new evidence - it formats what
`AiAgentService.converse()`/`compare_positions()` already produced (a
real answer grounded in real citations, per Phase 11's own system
prompt) into a real, shareable document. Lives alongside
`docx_writer.py`/`collection_export.py` (the other real python-docx-
touching code in this project), not under `application/` - all three
are real storage-adjacent adapters, not pure domain logic.
"""

from datetime import date
from pathlib import Path

from docx import Document


def build_answer_document(question: str, answer: str) -> Document:
    """Return a real, in-memory .docx `Document` for one real AI
    Assistant answer - a title (the question), the date, the real
    answer body (which already carries its own inline citations, per
    Phase 11's system prompt), and an honest disclaimer.

    Returns the `Document` object rather than saving it, so this stays
    testable without touching a real file - `export_answer_to_docx()`
    below is the thin save wrapper real callers use.
    """
    document = Document()
    document.add_heading(question, level=1)
    document.add_paragraph(date.today().isoformat())
    document.add_paragraph(answer)
    document.add_paragraph(
        "AI-generated research output, grounded in this library's real "
        "content - not a substitute for verifying the citations above "
        "against the real source, or for qualified scholarly guidance."
    )
    return document


def export_answer_to_docx(question: str, answer: str, output_path: Path) -> None:
    """Build and save a real AI answer export document to `output_path`."""
    build_answer_document(question, answer).save(output_path)
