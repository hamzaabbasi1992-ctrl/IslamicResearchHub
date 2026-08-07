"""Export real, AI-generated lecture-notes content (Phase 16 Milestone 2)
as a real .docx document.

Lives alongside `docx_writer.py`/`collection_export.py`/
`ai_answer_export.py`/`slide_deck_export.py` (the other real python-
docx/python-pptx-touching code in this project) rather than under
`application/` - all are real storage-adjacent adapters, not pure
domain logic.
"""

from pathlib import Path

from docx import Document

from islamic_research_hub.application.lecture_notes_extraction import ExtractedLectureSection


def build_lecture_notes_document(
    book_title: str, sections: tuple[ExtractedLectureSection, ...]
) -> Document:
    """Return a real, in-memory .docx `Document` - a title (the real
    source book), then one real heading + explanatory paragraph per
    `ExtractedLectureSection`, in the same order `sections` was given.

    Returns the `Document` object rather than saving it, so this stays
    testable without touching a real file -
    `export_lecture_notes_to_docx()` below is the thin save wrapper
    real callers use.
    """
    document = Document()
    document.add_heading(book_title, level=1)
    document.add_paragraph("AI-generated lecture notes, grounded in this library's real content")
    for section in sections:
        document.add_heading(section.heading, level=2)
        document.add_paragraph(section.content)
    return document


def export_lecture_notes_to_docx(
    book_title: str, sections: tuple[ExtractedLectureSection, ...], output_path: Path
) -> None:
    """Build and save a real lecture-notes export document to `output_path`."""
    build_lecture_notes_document(book_title, sections).save(output_path)
