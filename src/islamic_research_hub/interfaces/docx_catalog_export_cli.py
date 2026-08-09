"""CLI & Exporter to generate an alphabetical Word document (.docx) catalog of all books.

Includes ID, Title, Author, Maktaba (Library), Type (Category), Pages Count in separate table columns.
High performance XML stream generation for 100,000+ books.
"""

import argparse
import html
import logging
import sqlite3
from contextlib import closing
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor

LOGGER = logging.getLogger(__name__)

DEFAULT_DATABASE_PATH = Path("data/books.db")
DEFAULT_OUTPUT_PATH = Path("All_Books_Alphabetical_Catalog.docx")


def _xml_escape(text: str) -> str:
    """Sanitize text for valid Word XML insertion."""
    if not text:
        return ""
    # Strip invisible ASCII control codes except newlines/tabs
    sanitized = "".join(ch for ch in str(text) if ch in ("\n", "\r", "\t") or ord(ch) >= 32)
    return html.escape(sanitized)


def export_books_to_docx(database_path: Path, output_path: Path) -> tuple[int, Path]:
    """Export all books sorted alphabetically into a Word document with a formatted table."""
    if not database_path.is_file():
        raise FileNotFoundError(f"Database file not found: {database_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    with closing(sqlite3.connect(database_path)) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            """
            SELECT b.BookID,
                   COALESCE(b.Title, 'Untitled') AS Title,
                   COALESCE(b.Author, 'Unknown') AS Author,
                   COALESCE(l.Name, 'General') AS LibraryName,
                   COALESCE(b.Category, 'General') AS Category,
                   COALESCE(b.PageCount, 0) AS PageCount
            FROM Books b
            LEFT JOIN Libraries l ON l.LibraryID = b.LibraryID
            ORDER BY b.Title COLLATE NOCASE ASC
            """
        ).fetchall()

    doc = docx.Document()

    # Title Header
    header = doc.add_heading("Maktaba Master Catalog - Alphabetical Book List", level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary_p = doc.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_p.add_run(f"Total Books: {len(rows):,} | Sorted Alphabetically by Title")
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacing

    # Table Setup: 6 Columns (ID, Title, Author, Maktaba, Type, Pages Count)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ["ID", "Title", "Author", "Maktaba", "Type / Category", "Pages Count"]
    for i, h_text in enumerate(headers):
        cell_p = hdr_cells[i].paragraphs[0]
        cell_run = cell_p.add_run(h_text)
        cell_run.bold = True
        cell_run.font.size = Pt(10)
        cell_run.font.color.rgb = RGBColor(0, 51, 102)

    tbl_elm = table._tbl
    w_ns = nsdecls("w")

    # Ultra-fast XML row generation for 100,000+ items
    for row in rows:
        book_id = _xml_escape(str(row["BookID"]))
        title = _xml_escape(str(row["Title"]))
        author = _xml_escape(str(row["Author"]))
        maktaba = _xml_escape(str(row["LibraryName"]))
        category = _xml_escape(str(row["Category"]))
        page_count = _xml_escape(str(row["PageCount"]) if row["PageCount"] > 0 else "N/A")

        tr_xml = f"""
        <w:tr {w_ns}>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{book_id}</w:t></w:r></w:p></w:tc>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{title}</w:t></w:r></w:p></w:tc>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{author}</w:t></w:r></w:p></w:tc>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{maktaba}</w:t></w:r></w:p></w:tc>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{category}</w:t></w:r></w:p></w:tc>
          <w:tc><w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>{page_count}</w:t></w:r></w:p></w:tc>
        </w:tr>
        """
        tbl_elm.append(parse_xml(tr_xml))

    doc.save(output_path)
    return len(rows), output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Export alphabetical book catalog to Word document (.docx)")
    parser.add_argument("--database", type=Path, default=DEFAULT_DATABASE_PATH)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_PATH)
    args = parser.parse_args()

    count, path = export_books_to_docx(args.database, args.output)
    print(f"Successfully exported {count:,} books alphabetically to {path.resolve()}")


if __name__ == "__main__":
    main()
