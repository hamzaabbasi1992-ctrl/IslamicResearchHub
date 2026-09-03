import sys
sys.path.insert(0, r"F:\ISLAMIC RESEARCH HUB AI\src")
import os
import sqlite3
import json
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = r"F:\ISLAMIC RESEARCH HUB AI\data\books.db"
BASE_WAQIAT = r"F:\ISLAMIC RESEARCH HUB AI\WAQIAT ENCYCLOPEDIA FROM CLAUDE DESKTOP"
MASTER_TARGET_DIR = os.path.join(BASE_WAQIAT, "مکمل جلدیں واقعات کتابوں سے")

def clean_xml_text(s):
    if not s: return ""
    s = re.sub(r'</?[a-zA-Z][a-zA-Z0-9]*>', '', s)  # strip OCR markup tags: <urh1>, <urh2>, <urh3>, <qr>, <ar>, <arverse>, <hd>
    s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', s)
    return s

def add_simple_story_text(doc, text):
    paragraphs = text.split('\n')
    for p_text in paragraphs:
        p_text = clean_xml_text(p_text.strip())
        if not p_text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.left_indent = Inches(0.2)
        
        run = p.add_run(p_text)
        run.font.name = 'Jameel Noori Nastaleeq'
        run.font.size = Pt(13.5)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_p.paragraph_format.space_before = Pt(6)
    sep_p.paragraph_format.space_after = Pt(12)
    s_run = sep_p.add_run("❖ ❖ ❖")
    s_run.font.name = 'Jameel Noori Nastaleeq'
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def sync_all():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute("""
        SELECT ec.EventCandidateID, ec.BookID, b.Title AS BookTitle, ec.Title, ec.ChunkStartPage, ec.ExtractedDataJson
        FROM EventCandidates ec
        JOIN Books b ON ec.BookID = b.BookID
        WHERE ec.Status='confirmed'
        ORDER BY ec.BookID, ec.EventCandidateID
    """)
    rows = cur.fetchall()

    print(f"=======================================================")
    print(f" SYNCING WAQIAT DATABASE: {len(rows)} TOTAL WAQIAT FOUND")
    print(f"=======================================================")

    items = []
    for r in rows:
        data = json.loads(r[5]) if r[5] else {}
        items.append({
            'id': r[0],
            'book_id': r[1],
            'book_title': r[2],
            'title': clean_xml_text(r[3]),
            'page': r[4],
            'subject': clean_xml_text(data.get('subject') or 'اخلاق و موعظت'),
            'key_figures': data.get('key_figures') or [],
            'text': clean_xml_text(data.get('quoted_excerpt') or data.get('background') or ''),
            'citation': clean_xml_text(data.get('citation') or f"{r[2]}، صفحہ {r[4]}")
        })

    # Save data.json & data.js
    json_str = json.dumps(items, ensure_ascii=False)
    js_content = f"window.WAQIAT_DATABASE = {json_str};\n"

    target_folders = [
        os.path.join(BASE_WAQIAT, "SEARCH APP"),
        os.path.join(BASE_WAQIAT, "RESEARCH APP"),
        r"F:\ISLAMIC RESEARCH HUB AI\waqiat_dashboard",
        r"F:\ISLAMIC RESEARCH HUB AI\RESEARCH APP",
        r"F:\ISLAMIC RESEARCH HUB AI\mobile\app\src\main\assets"
    ]

    for folder in target_folders:
        os.makedirs(folder, exist_ok=True)
        with open(os.path.join(folder, "data.js"), "w", encoding="utf-8") as f:
            f.write(js_content)
        with open(os.path.join(folder, "waqiat_database.json"), "w", encoding="utf-8") as f:
            f.write(json_str)
        print(f"Updated App Folder: {folder}")

    # Rebuild Grand Master Word Document
    os.makedirs(MASTER_TARGET_DIR, exist_ok=True)
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36); title_p.paragraph_format.space_after = Pt(8)
    r = title_p.add_run("واقعات انسائیکلوپیڈیا")
    r.font.name = 'Jameel Noori Nastaleeq'; r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub_p.paragraph_format.space_after = Pt(14)
    r2 = sub_p.add_run(f"گرینڈ ماسٹر انسائیکلوپیڈیا — تمام کتب و سیریز (مکمل {len(rows)} مستند واقعات)")
    r2.font.name = 'Jameel Noori Nastaleeq'; r2.font.size = Pt(24); r2.font.bold = True; r2.font.color.rgb = RGBColor(0x4A, 0x77, 0x7A)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER; meta_p.paragraph_format.space_after = Pt(36)
    rm = meta_p.add_run(f"مجموعی مستند واقعات: {len(rows)} | جامع انسائیکلوپیڈیا برائے مطالعہ و موعظت | اسلامک ریسرچ ہب")
    rm.font.name = 'Jameel Noori Nastaleeq'; rm.font.size = Pt(14); rm.font.italic = True; rm.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    current_b_id = None
    g_count = 0

    for cid, b_id, b_title, title, sp, data_json in rows:
        g_count += 1
        if b_id != current_b_id:
            current_b_id = b_id
            v_head = doc.add_paragraph()
            v_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            v_head.paragraph_format.space_before = Pt(28); v_head.paragraph_format.space_after = Pt(12)
            vr = v_head.add_run(clean_xml_text(f"❖ ماخذ: {b_title} ❖"))
            vr.font.name = 'Jameel Noori Nastaleeq'; vr.font.size = Pt(22); vr.font.bold = True; vr.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        data = json.loads(data_json) if data_json else {}
        matn = data.get('quoted_excerpt') or data.get('background') or ''
        cit = data.get('citation') or f"{b_title}، صفحہ {sp}"
        subj = data.get('subject') or 'اخلاق و موعظت'
        figs = ", ".join(data.get('key_figures') or [])

        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; hp.paragraph_format.space_before = Pt(14); hp.paragraph_format.space_after = Pt(4)
        hrun = hp.add_run(clean_xml_text(f"واقعہ {g_count}: {title}"))
        hrun.font.name = 'Jameel Noori Nastaleeq'; hrun.font.size = Pt(17); hrun.font.bold = True; hrun.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; mp.paragraph_format.space_after = Pt(6)
        meta_str = f"موضوع: {subj}"
        if figs: meta_str += f" | شخصیت: {figs}"
        meta_str += f" | حوالہ: {cit}"
        mrun = mp.add_run(clean_xml_text(meta_str))
        mrun.font.name = 'Jameel Noori Nastaleeq'; mrun.font.size = Pt(11); mrun.font.italic = True; mrun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        add_simple_story_text(doc, matn)

    master_path = os.path.join(MASTER_TARGET_DIR, f"واقعات انسائیکلوپیڈیا — گرینڈ ماسٹر انسائیکلوپیڈیا (تمام کتب و سیریز — {len(rows)} واقعات).docx")

    # Remove superseded snapshots from earlier counts so old versions don't pile up
    for fname in os.listdir(MASTER_TARGET_DIR):
        if fname.startswith("واقعات انسائیکلوپیڈیا — گرینڈ ماسٹر انسائیکلوپیڈیا") and fname.endswith(".docx"):
            fpath = os.path.join(MASTER_TARGET_DIR, fname)
            if os.path.abspath(fpath) != os.path.abspath(master_path):
                try:
                    os.remove(fpath)
                except OSError:
                    pass

    try:
        doc.save(master_path)
    except PermissionError:
        doc.save(os.path.join(MASTER_TARGET_DIR, f"واقعات انسائیکلوپیڈیا — گرینڈ ماسٹر انسائیکلوپیڈیا (تمام کتب و سیریز — {len(rows)} واقعات)_جدید.docx"))

    conn.close()
    print(f"SUCCESS: All Web App Folders & Grand Master Word Document synced to {len(rows)} Waqiat!")

if __name__ == "__main__":
    sync_all()
