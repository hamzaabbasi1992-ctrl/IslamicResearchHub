"""Tests for docx catalog export functionality."""

import sqlite3
from contextlib import closing
from pathlib import Path

import docx
import pytest

from islamic_research_hub.interfaces.docx_catalog_export_cli import export_books_to_docx


def _seed_test_db(db_path: Path) -> None:
    with closing(sqlite3.connect(db_path)) as conn:
        conn.executescript(
            """
            CREATE TABLE Libraries (
                LibraryID INTEGER PRIMARY KEY,
                Name TEXT NOT NULL
            );
            CREATE TABLE Books (
                BookID INTEGER PRIMARY KEY,
                LibraryID INTEGER,
                Title TEXT NOT NULL,
                Author TEXT,
                Publisher TEXT,
                Language TEXT,
                Category TEXT,
                PageCount INTEGER
            );
            INSERT INTO Libraries VALUES (1, 'Main Library');
            INSERT INTO Books VALUES (101, 1, 'Zahr al-Ruba', 'Imam Suyuti', 'Dar', 'Arabic', 'Hadith', 150);
            INSERT INTO Books VALUES (102, 1, 'Al-Bidayah wa al-Nihayah', 'Ibn Kathir', 'Dar', 'Arabic', 'History', 500);
            """
        )


def test_export_books_to_docx(tmp_path: Path) -> None:
    db_path = tmp_path / "books.db"
    _seed_test_db(db_path)
    out_docx = tmp_path / "catalog.docx"

    count, path = export_books_to_docx(db_path, out_docx)
    assert count == 2
    assert path.is_file()

    doc = docx.Document(path)
    assert len(doc.tables) == 1
    table = doc.tables[0]

    # Check header
    headers = [cell.text for cell in table.rows[0].cells]
    assert headers == ["ID", "Title", "Author", "Maktaba", "Type / Category", "Pages Count"]

    # Check alphabetical ordering (Al-Bidayah comes before Zahr al-Ruba)
    first_row_title = table.rows[1].cells[1].text
    assert "Al-Bidayah" in first_row_title
