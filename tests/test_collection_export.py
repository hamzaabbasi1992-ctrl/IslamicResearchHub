"""Tests for building a real .docx export document for one collection."""

from pathlib import Path

from docx import Document

from islamic_research_hub.research_notes.collection_export import (
    CollectionExportItem,
    build_export_document,
    export_collection_to_docx,
)


def test_build_export_document_includes_the_collection_name_as_a_heading() -> None:
    document = build_export_document("Zakat research", items=())

    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "Zakat research" in headings


def test_build_export_document_with_no_items_says_so_honestly() -> None:
    document = build_export_document("Empty collection", items=())

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "no items yet" in all_text.lower()


def test_build_export_document_includes_each_items_real_content_and_citation() -> None:
    items = (
        CollectionExportItem(
            book_title="Book of Fiqh",
            author="Author One",
            volume_number=None,
            page_number=5,
            content="Real page content about zakat.",
        ),
        CollectionExportItem(
            book_title="Book of Hadith",
            author="Author Two",
            volume_number=2,
            page_number=10,
            content="Real page content about sabr.",
        ),
    )

    document = build_export_document("Zakat research", items)

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "Book of Fiqh" in all_text
    assert "Author One" in all_text
    assert "Page 5" in all_text
    assert "Real page content about zakat." in all_text
    assert "Book of Hadith" in all_text
    assert "Volume 2" in all_text
    assert "Real page content about sabr." in all_text


def test_export_collection_to_docx_writes_a_real_readable_file(tmp_path: Path) -> None:
    items = (
        CollectionExportItem(
            book_title="Book of Fiqh",
            author="Author One",
            volume_number=None,
            page_number=5,
            content="Real page content.",
        ),
    )
    output_path = tmp_path / "Zakat research.docx"

    export_collection_to_docx("Zakat research", items, output_path)

    assert output_path.is_file()
    reopened = Document(output_path)
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "Book of Fiqh" in all_text
