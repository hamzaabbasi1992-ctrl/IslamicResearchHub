"""Tests for the local .docx storage adapter - real python-docx round-trips, no Qt."""

from pathlib import Path

import pytest
from docx import Document

from islamic_research_hub.research_notes.docx_writer import (
    LocalDocxStorage,
    NoteFileLockedError,
)
from islamic_research_hub.research_notes.research_notes_manager import Quotation


def _quotation(**overrides: object) -> Quotation:
    defaults: dict[object, object] = dict(
        book_title="Book of Fiqh",
        author="Author One",
        volume=2,
        chapter="Chapter of Prayer",
        page_number=17,
        selected_text="The rules of jurisprudence are extensive.",
    )
    defaults.update(overrides)
    return Quotation(**defaults)  # type: ignore[arg-type]


def _storage(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> LocalDocxStorage:
    """A LocalDocxStorage whose real notes folder is redirected under
    tmp_path, so tests never touch the real Documents folder."""
    storage = LocalDocxStorage()
    monkeypatch.setattr(storage, "notes_folder", lambda: tmp_path)
    return storage


def _paragraph_texts(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def test_create_document_creates_a_real_empty_docx(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)

    path = storage.create_document("Thesis")

    assert path.is_file()
    assert path.name == "Thesis.docx"
    assert _paragraph_texts(path) == []


def test_create_document_sanitizes_illegal_filename_characters(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)

    path = storage.create_document('Fiqh: Q&A? "Notes"')

    assert path.is_file()
    assert not any(char in path.name for char in '\\/:*?"<>|')


def test_create_document_never_overwrites_an_existing_name(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    first = storage.create_document("Notes")
    storage.append_quotation(first, _quotation())

    second = storage.create_document("Notes")

    assert second != first
    assert second.is_file()
    # The first document's real content survived - never silently overwritten.
    assert "The rules of jurisprudence are extensive." in _paragraph_texts(first)


def test_list_documents_returns_every_real_docx_alphabetically(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    storage.create_document("Zebra")
    storage.create_document("Alpha")

    documents = storage.list_documents()

    assert [path.stem for path in documents] == ["Alpha", "Zebra"]


def test_append_quotation_includes_every_real_field(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")

    storage.append_quotation(path, _quotation())

    text = "\n".join(_paragraph_texts(path))
    assert "Book of Fiqh" in text
    assert "Author One" in text
    assert "2" in text  # volume
    assert "Chapter of Prayer" in text
    assert "17" in text  # page
    assert "The rules of jurisprudence are extensive." in text
    assert "My Notes" in text


def test_append_quotation_omits_missing_optional_fields(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Volume/chapter are real, sometimes-absent facts - not fabricated
    when unknown, matching format_citation()'s own precedent."""
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")

    storage.append_quotation(
        path, _quotation(volume=None, chapter=None, author=None)
    )

    text = "\n".join(_paragraph_texts(path))
    assert "Volume:" not in text
    assert "Chapter:" not in text
    assert "Unknown" in text  # author falls back to a real, honest placeholder


def test_append_quotation_appends_after_existing_content_never_overwrites(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")

    storage.append_quotation(path, _quotation(selected_text="First real quotation"))
    storage.append_quotation(path, _quotation(selected_text="Second real quotation"))

    text = "\n".join(_paragraph_texts(path))
    assert "First real quotation" in text
    assert "Second real quotation" in text


def test_append_quotation_translates_a_real_permission_error(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A document open in Word holds a real exclusive lock on Windows -
    `Document.save()` then raises `PermissionError`. This must become a
    real, catchable `NoteFileLockedError` with a friendly message, never
    an uncaught crash. A fake Document stands in for the real save call,
    since reproducing an actual cross-process Word file lock isn't
    something an automated test can do reliably."""
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")

    class _LockedDocument:
        def add_paragraph(self, _text: str) -> None:
            pass

        def save(self, _path: Path) -> None:
            raise PermissionError("locked by another process")

    monkeypatch.setattr(
        "islamic_research_hub.research_notes.docx_writer.Document",
        lambda *_args, **_kwargs: _LockedDocument(),
    )

    with pytest.raises(NoteFileLockedError, match="currently open"):
        storage.append_quotation(path, _quotation())


def test_find_documents_mentioning_finds_a_real_saved_quotation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")
    storage.append_quotation(path, _quotation(book_title="Book of Fiqh"))

    matches = storage.find_documents_mentioning("Book of Fiqh")

    assert matches == (path,)


def test_find_documents_mentioning_excludes_documents_about_other_books(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    path = storage.create_document("Notes")
    storage.append_quotation(path, _quotation(book_title="Book of Fiqh"))

    matches = storage.find_documents_mentioning("A Completely Different Book")

    assert matches == ()


def test_find_documents_mentioning_checks_every_document(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    storage = _storage(tmp_path, monkeypatch)
    thesis = storage.create_document("Thesis")
    storage.append_quotation(thesis, _quotation(book_title="Book of Fiqh"))
    fiqh_notes = storage.create_document("Fiqh Research")
    storage.append_quotation(fiqh_notes, _quotation(book_title="Book of Fiqh"))
    unrelated = storage.create_document("Unrelated")
    storage.append_quotation(unrelated, _quotation(book_title="Some Other Book"))

    matches = storage.find_documents_mentioning("Book of Fiqh")

    assert set(matches) == {thesis, fiqh_notes}
