"""Tests for building a real .docx export document for AI-generated lecture notes."""

from pathlib import Path

from docx import Document

from islamic_research_hub.application.lecture_notes_extraction import ExtractedLectureSection
from islamic_research_hub.research_notes.lecture_notes_export import (
    build_lecture_notes_document,
    export_lecture_notes_to_docx,
)

_SECTIONS = (
    ExtractedLectureSection(
        heading="The Battle of Badr",
        content="Fought in 2 AH, a decisive early victory for the Muslims.",
    ),
    ExtractedLectureSection(
        heading="The Treaty of Hudaybiyyah",
        content="Signed in 6 AH, a real turning point despite its apparent terms.",
    ),
)


def test_build_lecture_notes_document_includes_a_title_naming_the_book() -> None:
    document = build_lecture_notes_document("Seerah of the Prophet", _SECTIONS)

    assert document.paragraphs[0].text == "Seerah of the Prophet"


def test_build_lecture_notes_document_includes_every_real_section_heading_and_content() -> None:
    document = build_lecture_notes_document("Seerah of the Prophet", _SECTIONS)

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "The Battle of Badr" in all_text
    assert "Fought in 2 AH, a decisive early victory for the Muslims." in all_text
    assert "The Treaty of Hudaybiyyah" in all_text
    assert "Signed in 6 AH, a real turning point despite its apparent terms." in all_text


def test_build_lecture_notes_document_with_no_sections_still_has_a_title() -> None:
    document = build_lecture_notes_document("An Empty Book", ())

    assert document.paragraphs[0].text == "An Empty Book"


def test_export_lecture_notes_to_docx_writes_a_real_readable_file(tmp_path: Path) -> None:
    output_path = tmp_path / "notes.docx"

    export_lecture_notes_to_docx("Seerah of the Prophet", _SECTIONS, output_path)

    assert output_path.is_file()
    reopened = Document(output_path)
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "The Battle of Badr" in all_text
