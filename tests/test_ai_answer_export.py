"""Tests for building a real .docx export document for one AI Assistant answer."""

from pathlib import Path

from docx import Document

from islamic_research_hub.research_notes.ai_answer_export import (
    build_answer_document,
    export_answer_to_docx,
)


def test_build_answer_document_includes_the_question_as_a_heading() -> None:
    document = build_answer_document("What does this library say about patience?", "A real answer.")

    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "What does this library say about patience?" in headings


def test_build_answer_document_includes_the_real_answer_text() -> None:
    document = build_answer_document(
        "A question.", "A real grounded answer with a real citation (Book of Fiqh, Page 5)."
    )

    all_text = "\n".join(p.text for p in document.paragraphs)
    assert "A real grounded answer with a real citation (Book of Fiqh, Page 5)." in all_text


def test_build_answer_document_includes_an_honest_disclaimer() -> None:
    document = build_answer_document("A question.", "An answer.")

    all_text = "\n".join(p.text for p in document.paragraphs).lower()
    assert "not a substitute" in all_text


def test_export_answer_to_docx_writes_a_real_readable_file(tmp_path: Path) -> None:
    output_path = tmp_path / "answer.docx"

    export_answer_to_docx("A real question", "A real answer.", output_path)

    assert output_path.is_file()
    reopened = Document(output_path)
    all_text = "\n".join(p.text for p in reopened.paragraphs)
    assert "A real answer." in all_text
