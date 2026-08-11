"""Tests for the MCP server's tools (search, citation, health check,
bookmarks, collections, notes, citation candidates, exports)."""

import asyncio
import json
import sqlite3
from pathlib import Path

import pytest
from docx import Document
from mcp.server.mcpserver.exceptions import ToolError

from islamic_research_hub.domain.models.book import Book, Page
from islamic_research_hub.infrastructure.persistence.master_book_repository import (
    MasterBookRepository,
)
from islamic_research_hub.infrastructure.persistence.migration_runner import MigrationRunner
from islamic_research_hub.interfaces.mcp_server.server import build_server
from islamic_research_hub.research_notes.docx_writer import LocalDocxStorage


def _seed_database(database_path: Path) -> None:
    """Import one book with searchable content into a fresh master database."""
    book = Book(
        information={"Name": "Book of Fiqh", "ANAME": "Author One"},
        categories=(),
        table_of_contents=(),
        pages=(Page(1, 1, "The rules of jurisprudence in fiqh are extensive", "Plain"),),
    )
    MasterBookRepository().import_books(
        database_path, (book,), (database_path.parent / "source.mjbz",)
    )


def _migrated_database(tmp_path: Path) -> Path:
    """Create a real, fully-migrated database with one real book (BookID 1)."""
    database_path = tmp_path / "books.db"
    book = Book(
        information={"Name": "Book of Fiqh", "ANAME": "Author One"},
        categories=(),
        table_of_contents=(),
        pages=(
            Page(1, 1, "The rules of jurisprudence in fiqh are extensive", "Plain"),
            Page(2, 2, "More content on a second page", "Plain"),
        ),
    )
    MasterBookRepository().import_books(
        database_path, (book,), (database_path.parent / "source.mjbz",)
    )
    with sqlite3.connect(database_path) as connection:
        MigrationRunner().migrate(connection)
    return database_path


def _call(server, tool_name: str, arguments: dict):
    """Call an MCP tool synchronously and return its structured result."""
    result = asyncio.run(server.call_tool(tool_name, arguments))
    if result.structured_content is not None:
        return result.structured_content["result"]
    return json.loads(result.content[0].text)


def test_search_text_returns_ranked_matches(tmp_path: Path) -> None:
    """search_text returns the same match data SqliteBookSearchRepository would."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    results = _call(server, "search_text", {"query": "jurisprudence"})

    assert len(results) == 1
    assert results[0]["title"] == "Book of Fiqh"
    assert results[0]["author"] == "Author One"
    assert results[0]["page_number"] == 1
    assert "jurisprudence" in results[0]["excerpt"].lower()


def test_search_text_rejects_blank_query(tmp_path: Path) -> None:
    """A blank query surfaces BookSearchService's own validation error."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    with pytest.raises(ToolError, match="Search query must not be empty"):
        asyncio.run(server.call_tool("search_text", {"query": "  "}))


def test_get_citation_returns_formatted_string(tmp_path: Path) -> None:
    """get_citation formats book/page/paragraph without a volume number."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    citation = _call(
        server, "get_citation", {"book_id": 1, "page_number": 1, "paragraph_index": 3}
    )

    assert citation == "Book Book of Fiqh, Page 1, Paragraph 3"


def test_get_open_link_returns_a_maktaba_link(tmp_path: Path) -> None:
    """get_open_link returns the same format build_maktaba_link() produces."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    link = _call(server, "get_open_link", {"book_id": 1, "page_number": 3})

    assert link == "maktaba://open?book=1&page=3"


def test_get_citation_raises_for_unknown_book(tmp_path: Path) -> None:
    """An unknown book_id raises a clear error rather than a blank citation."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    with pytest.raises(ToolError, match="No book found with book_id=999"):
        asyncio.run(
            server.call_tool(
                "get_citation", {"book_id": 999, "page_number": 1, "paragraph_index": 1}
            )
        )


def test_health_check_reports_reachable_database(tmp_path: Path) -> None:
    """health_check confirms the configured database file exists."""
    database_path = tmp_path / "books.db"
    _seed_database(database_path)
    server = build_server(database_path)

    status = _call(server, "health_check", {})

    assert status["reachable"] is True
    assert status["database_path"] == str(database_path)


def test_health_check_reports_missing_database(tmp_path: Path) -> None:
    """health_check reports unreachable for a path with no database file yet."""
    database_path = tmp_path / "missing.db"
    server = build_server(database_path)

    status = _call(server, "health_check", {})

    assert status["reachable"] is False


# --- Bookmarks ---


def test_add_bookmark_then_list_then_remove(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))

    _call(server, "add_bookmark", {"book_id": 1, "page_number": 2})
    assert _call(server, "list_bookmarked_pages", {"book_id": 1}) == [2]

    recent = _call(server, "list_recent_bookmarks", {"limit": 5})
    assert recent == [{"book_id": 1, "page_number": 2, "title": "Book of Fiqh"}]

    _call(server, "remove_bookmark", {"book_id": 1, "page_number": 2})
    assert _call(server, "list_bookmarked_pages", {"book_id": 1}) == []


# --- Collections ---


def test_create_collection_add_item_and_list(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))

    created = _call(server, "create_collection", {"name": "Zakat research"})
    collection_id = created["collection_id"]

    collections = _call(server, "list_collections", {})
    assert collections[0]["name"] == "Zakat research"
    assert collections[0]["item_count"] == 0

    _call(server, "add_to_collection", {"collection_id": collection_id, "book_id": 1, "page_number": 1})
    items = _call(server, "list_collection_items", {"collection_id": collection_id})
    assert len(items) == 1
    assert items[0]["book_title"] == "Book of Fiqh"

    _call(
        server,
        "remove_from_collection",
        {"collection_id": collection_id, "book_id": 1, "page_number": 1},
    )
    assert _call(server, "list_collection_items", {"collection_id": collection_id}) == []


def test_create_collection_with_taken_name_raises(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))
    _call(server, "create_collection", {"name": "Zakat research"})

    with pytest.raises(ToolError, match="already exists"):
        asyncio.run(server.call_tool("create_collection", {"name": "Zakat research"}))


def test_delete_collection_removes_it(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))
    created = _call(server, "create_collection", {"name": "Zakat research"})

    _call(server, "delete_collection", {"collection_id": created["collection_id"]})

    assert _call(server, "list_collections", {}) == []


# --- Research notes ---


def test_create_note_document_then_save_quotation_then_find(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(LocalDocxStorage, "notes_folder", lambda self: tmp_path)
    server = build_server(_migrated_database(tmp_path))

    document_path = _call(server, "create_note_document", {"name": "Thesis"})
    assert Path(document_path).is_file()

    saved = _call(
        server,
        "save_quotation",
        {
            "document_path": document_path,
            "book_title": "Book of Fiqh",
            "selected_text": "The rules of jurisprudence are extensive.",
            "author": "Author One",
            "page_number": 1,
        },
    )
    assert saved["saved"] is True

    matches = _call(server, "find_notes_mentioning_book", {"book_title": "Book of Fiqh"})
    assert matches == [document_path]

    document = Document(document_path)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    assert "The rules of jurisprudence are extensive." in paragraph_texts


# --- Citation candidates ---


def _insert_citation_candidate(database_path: Path) -> None:
    with sqlite3.connect(database_path) as connection:
        connection.execute(
            "CREATE TABLE IF NOT EXISTS CitationCandidates ("
            "CitingBookID INTEGER NOT NULL, CitingPageNo INTEGER NOT NULL, "
            "CitingParagraphID INTEGER, CitedBookID INTEGER NOT NULL, "
            "MatchedTitleText TEXT NOT NULL, MatchType TEXT NOT NULL, "
            "Status TEXT NOT NULL DEFAULT 'pending', "
            "PRIMARY KEY (CitingBookID, CitingPageNo, CitedBookID))"
        )
        connection.execute(
            "INSERT INTO CitationCandidates "
            "(CitingBookID, CitingPageNo, CitingParagraphID, CitedBookID, MatchedTitleText, MatchType) "
            "VALUES (1, 1, NULL, 1, 'Book of Fiqh', 'exact')"
        )
        connection.commit()


def test_list_and_dismiss_citation_candidate(tmp_path: Path) -> None:
    database_path = _migrated_database(tmp_path)
    _insert_citation_candidate(database_path)
    server = build_server(database_path)

    candidates = _call(server, "list_citation_candidates", {})
    assert len(candidates) == 1
    assert candidates[0]["status"] == "pending"

    assert _call(server, "count_citation_candidates", {}) == 1

    _call(
        server,
        "dismiss_citation_candidate",
        {"citing_book_id": 1, "citing_page_no": 1, "cited_book_id": 1},
    )
    assert _call(server, "list_citation_candidates", {}) == []
    assert _call(server, "list_citation_candidates", {"include_dismissed": True})[0]["status"] == "dismissed"


# --- Exports ---


def test_export_answer_to_docx(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))
    output_path = tmp_path / "answer.docx"

    _call(
        server,
        "export_answer_to_docx",
        {"question": "What is zakat?", "answer": "Zakat is...", "output_path": str(output_path)},
    )

    assert output_path.is_file()
    texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "Zakat is..." in texts


def test_export_collection_to_docx(tmp_path: Path) -> None:
    database_path = _migrated_database(tmp_path)
    server = build_server(database_path)
    created = _call(server, "create_collection", {"name": "Fiqh notes"})
    _call(
        server,
        "add_to_collection",
        {"collection_id": created["collection_id"], "book_id": 1, "page_number": 1},
    )
    output_path = tmp_path / "collection.docx"

    result = _call(
        server,
        "export_collection_to_docx",
        {"collection_id": created["collection_id"], "output_path": str(output_path)},
    )

    assert result["item_count"] == 1
    assert output_path.is_file()
    texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert any("jurisprudence" in text for text in texts)


def test_export_article_to_docx(tmp_path: Path) -> None:
    server = build_server(_migrated_database(tmp_path))
    output_path = tmp_path / "article.docx"

    result = _call(
        server,
        "export_article_to_docx",
        {
            "title": "On Zakat",
            "sections": [{"heading": "Introduction", "body": "Zakat is a pillar of Islam."}],
            "sources": ["Book Book of Fiqh, Page 1, Paragraph 1"],
            "output_path": str(output_path),
        },
    )

    assert result["section_count"] == 1
    assert output_path.is_file()
    texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "Zakat is a pillar of Islam." in texts
    assert any("Works Cited" in text for text in texts)
